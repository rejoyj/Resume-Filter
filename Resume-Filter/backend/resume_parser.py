import re
import spacy
import pandas as pd
import PyPDF2
import docx
from datetime import datetime
import json
import os
from collections import defaultdict
import phonenumbers
from phonenumbers import geocoder, carrier
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.chunk import ne_chunk
from nltk.tag import pos_tag
import logging
from typing import List, Dict, Optional, Any

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

nlp = spacy.load("en_core_web_sm")

class ResumeParser:
    def __init__(self):
        """Initialize the Resume Parser with all required dependencies"""
        self._setup_nltk_data()
        self._setup_spacy_model()
        self._initialize_skill_keywords()
        self._initialize_education_patterns()
        
        # Statistics tracking
        self.processing_stats = {
            'total_processed': 0,
            'successful_extractions': defaultdict(int),
            'failed_files': [],
            'processing_time': 0
        }

    def _setup_nltk_data(self):
        """Download and setup required NLTK data"""
        nltk_requirements = [
            'punkt', 'stopwords', 'averaged_perceptron_tagger',
            'maxent_ne_chunker', 'words'
        ]
        
        for requirement in nltk_requirements:
            try:
                nltk.data.find(f'tokenizers/{requirement}' if requirement == 'punkt' 
                              else f'corpora/{requirement}' if requirement in ['stopwords', 'words']
                              else f'taggers/{requirement}' if 'tagger' in requirement
                              else f'chunkers/{requirement}')
            except LookupError:
                logger.info(f"Downloading NLTK data: {requirement}")
                nltk.download(requirement, quiet=True)

    def _setup_spacy_model(self):
        """Setup spaCy model with fallback"""
        try:
            self.nlp = spacy.load("en_core_web_sm")
            logger.info("✓ spaCy model 'en_core_web_sm' loaded successfully")
        except OSError:
            logger.warning("⚠️  spaCy model 'en_core_web_sm' not found. Using fallback parsing.")
            logger.warning("Install it using: python -m spacy download en_core_web_sm")
            self.nlp = None

    def _initialize_skill_keywords(self):
        """Initialize comprehensive skills keywords database"""
        self.skills_keywords = {
            'programming_languages': [
                'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby', 
                'go', 'rust', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'bash', 
                'powershell', 'html', 'css', 'sql', 'nosql', 'c', 'assembly', 'vba',
                'dart', 'elixir', 'haskell', 'lua', 'objective-c', 'groovy', 'clojure'
            ],
            'frameworks_libraries': [
                'django', 'flask', 'fastapi', 'react', 'angular', 'vue', 'svelte', 'spring',
                'express', 'laravel', 'rails', 'asp.net', 'bootstrap', 'jquery', 'nodejs',
                'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'opencv', 'pandas', 'numpy',
                'matplotlib', 'seaborn', 'plotly', 'streamlit', 'gradio', 'huggingface',
                'next.js', 'nuxt.js', 'gatsby', 'redux', 'mobx', 'webpack', 'vite'
            ],
            'databases': [
                'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle', 'cassandra',
                'elasticsearch', 'neo4j', 'dynamodb', 'firebase', 'supabase', 'mariadb',
                'couchdb', 'influxdb', 'clickhouse', 'snowflake', 'bigquery'
            ],
            'cloud_devops': [
                'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'travis',
                'ansible', 'terraform', 'vagrant', 'helm', 'istio', 'prometheus',
                'grafana', 'elk', 'splunk', 'datadog', 'newrelic', 'ci/cd', 'gitlab-ci',
                'github-actions', 'circleci', 'heroku', 'vercel', 'netlify'
            ],
            'tools_technologies': [
                'git', 'github', 'gitlab', 'bitbucket', 'linux', 'windows', 'macos',
                'jira', 'confluence', 'slack', 'teams', 'trello', 'asana', 'notion',
                'figma', 'sketch', 'adobe', 'photoshop', 'illustrator', 'xd',
                'postman', 'insomnia', 'swagger', 'apache', 'nginx', 'elasticsearch'
            ],
            'concepts_methodologies': [
                'machine learning', 'artificial intelligence', 'deep learning', 'data science',
                'web development', 'mobile development', 'devops', 'cloud computing',
                'microservices', 'api', 'rest', 'graphql', 'agile', 'scrum', 'kanban',
                'testing', 'automation', 'security', 'networking', 'blockchain',
                'cybersecurity', 'data analysis', 'business intelligence', 'etl',
                'big data', 'nlp', 'computer vision', 'iot', 'edge computing'
            ],
            'soft_skills': [
                'leadership', 'communication', 'teamwork', 'problem solving', 'analytical',
                'critical thinking', 'creativity', 'adaptability', 'time management',
                'project management', 'collaboration', 'mentoring', 'presentation',
                'negotiation', 'strategic thinking', 'innovation', 'customer service'
            ]
        }
        
        # Flatten all skills for easy searching
        self.all_skills = []
        for category in self.skills_keywords.values():
            self.all_skills.extend(category)

    def _initialize_education_patterns(self):
        """Initialize education keywords and patterns"""
        self.education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'diploma', 'certificate',
            'b.e.', 'b.tech', 'b.sc', 'b.com', 'b.a.', 'm.e.', 'm.tech', 'm.sc',
            'm.com', 'm.a.', 'mba', 'bca', 'mca', 'engineering', 'computer science',
            'information technology', 'electronics', 'mechanical', 'civil', 'electrical',
            'degree', 'graduation', 'post graduation', 'undergraduate', 'graduate'
        ]
        
        self.degree_patterns = [
            r'b\.?e\.?\s*(?:in\s+)?([^,\n.]+)',
            r'b\.?tech\.?\s*(?:in\s+)?([^,\n.]+)',
            r'm\.?e\.?\s*(?:in\s+)?([^,\n.]+)',
            r'm\.?tech\.?\s*(?:in\s+)?([^,\n.]+)',
            r'bachelor.*?(?:of|in)\s+([^,\n.]+)',
            r'master.*?(?:of|in)\s+([^,\n.]+)',
            r'mba\s*(?:in\s+)?([^,\n.]*)',
            r'b\.?sc\.?\s*(?:in\s+)?([^,\n.]+)',
            r'm\.?sc\.?\s*(?:in\s+)?([^,\n.]+)',
            r'phd\s*(?:in\s+)?([^,\n.]*)',
            r'doctorate\s*(?:in\s+)?([^,\n.]*)'
        ]

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file with enhanced error handling"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    logger.warning(f"PDF {file_path} is encrypted. Attempting to decrypt...")
                    try:
                        pdf_reader.decrypt('')
                    except:
                        logger.error(f"Could not decrypt PDF {file_path}")
                        return ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1} of {file_path}: {str(e)}")
                        continue
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file with enhanced error handling"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            return ""

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file with multiple encoding support"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read().strip()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error reading TXT {file_path}: {str(e)}")
                return ""
        
        logger.error(f"Could not decode TXT file {file_path} with any encoding")
        return ""

    def extract_name(self, text: str) -> Optional[str]:
        """Extract name from resume text with improved accuracy"""
        lines = text.split('\n')
        
        # Strategy 1: Look for name in first few lines
        for i, line in enumerate(lines[:5]):
            line = line.strip()
            if len(line) > 2 and len(line) < 60:
                words = line.split()
                if 2 <= len(words) <= 4:
                    # Check if all words are likely names (alphabetic + some punctuation)
                    if all(re.match(r'^[A-Za-z\.\,\s]+$', word) for word in words):
                        # Avoid lines that look like headers or titles
                        if not any(keyword in line.lower() for keyword in 
                                 ['resume', 'cv', 'curriculum', 'profile', 'contact', 'email', 'phone']):
                            return ' '.join(word.strip('.,') for word in words).title()
        
        # Strategy 2: Use spaCy if available
        if self.nlp:
            doc = self.nlp(text[:1000])  # Process first 1000 characters
            person_entities = [ent.text.strip() for ent in doc.ents if ent.label_ == "PERSON"]
            
            if person_entities:
                # Return the first person entity that looks like a full name
                for entity in person_entities:
                    words = entity.split()
                    if 2 <= len(words) <= 4 and len(entity) < 50:
                        return entity.title()
        
        # Strategy 3: Pattern matching for name formats
        name_patterns = [
            r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]*\.?)*\s+[A-Z][a-z]+)$',  # First Last format
            r'^([A-Z][a-z]+\s+[A-Z]\.\s+[A-Z][a-z]+)$',  # First M. Last format
            r'^([A-Z][a-z]+\s+[A-Z][a-z]+\s+[A-Z][a-z]+)$'  # First Middle Last format
        ]
        
        for line in lines[:10]:
            line = line.strip()
            for pattern in name_patterns:
                match = re.match(pattern, line)
                if match:
                    return match.group(1)
        
        return None

    def extract_email(self, text: str) -> Optional[str]:
        """Extract email from resume text with improved validation"""
        # Enhanced email pattern that handles more formats
        email_pattern = r'\b[A-Za-z0-9]([A-Za-z0-9._%-]*[A-Za-z0-9])?@[A-Za-z0-9]([A-Za-z0-9.-]*[A-Za-z0-9])?\.[A-Za-z]{2,}\b'
        emails = re.findall(email_pattern, text, re.IGNORECASE)
        
        if emails:
            # Reconstruct full email from tuple groups
            full_emails = []
            for email_parts in emails:
                if isinstance(email_parts, tuple):
                    # This handles the case where regex groups are captured
                    continue
                else:
                    full_emails.append(email_parts)
            
            # Simple pattern for complete emails
            simple_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            simple_emails = re.findall(simple_pattern, text, re.IGNORECASE)
            
            if simple_emails:
                # Filter out obviously fake or template emails
                valid_emails = [email for email in simple_emails 
                              if not any(fake in email.lower() for fake in 
                                       ['example', 'test', 'sample', 'dummy', 'placeholder'])]
                return valid_emails[0] if valid_emails else simple_emails[0]
        
        return None

    def extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from resume text with international support"""
        # Multiple phone patterns for different formats
        phone_patterns = [
            r'(\+91[-.\s]?)?[6789]\d{9}',  # Indian mobile numbers
            r'(\+1[-.\s]?)?\(?[2-9]\d{2}\)?[-.\s]?[2-9]\d{2}[-.\s]?\d{4}',  # US numbers
            r'(\+44[-.\s]?)?(?:0)?[1-9]\d{8,9}',  # UK numbers
            r'(\+\d{1,3}[-.\s]?)?\d{10,14}',  # Generic international
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # Standard format
            r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}',  # (xxx) xxx-xxxx format
            r'\b\d{10}\b'  # 10 digit numbers
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            if matches:
                phone = matches[0]
                if isinstance(phone, tuple):
                    phone = ''.join(phone)
                
                # Clean up the phone number
                clean_phone = re.sub(r'[^\d+]', '', phone)
                
                # Validate length
                if 10 <= len(clean_phone.replace('+', '')) <= 15:
                    return clean_phone
        
        return None

    def extract_skills(self, text: str) -> Optional[List[str]]:
        """Extract skills from resume text with enhanced matching"""
        text_lower = text.lower()
        found_skills = set()
        
        # Method 1: Direct keyword matching
        for skill in self.all_skills:
            skill_lower = skill.lower()
            
            # Exact match or word boundary match
            if skill_lower in text_lower:
                # Check if it's a whole word match for better accuracy
                pattern = r'\b' + re.escape(skill_lower) + r'\b'
                if re.search(pattern, text_lower):
                    found_skills.add(skill.title())
        
        # Method 2: Skills section parsing
        skills_sections = [
            'skills', 'technical skills', 'core competencies', 'expertise',
            'technologies', 'programming languages', 'tools', 'frameworks'
        ]
        
        for section in skills_sections:
            section_pattern = rf'{section}[:\-\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*[A-Z][^:\n]*:|$)'
            match = re.search(section_pattern, text, re.IGNORECASE | re.MULTILINE)
            
            if match:
                skills_text = match.group(1).lower()
                for skill in self.all_skills:
                    if skill.lower() in skills_text:
                        found_skills.add(skill.title())
        
        # Method 3: Bullet point parsing
        bullet_patterns = [r'•\s*([^\n]+)', r'▪\s*([^\n]+)', r'-\s*([^\n]+)', r'\*\s*([^\n]+)']
        
        for pattern in bullet_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                match_lower = match.lower()
                for skill in self.all_skills:
                    if skill.lower() in match_lower:
                        found_skills.add(skill.title())
        
        return list(found_skills) if found_skills else None

    def extract_education(self, text: str) -> Optional[List[str]]:
        """Extract education information with improved parsing"""
        education_info = set()
        text_lower = text.lower()
        
        # Method 1: Degree pattern matching
        for pattern in self.degree_patterns:
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                if match.strip() and len(match.strip()) > 2:
                    education_info.add(match.strip().title())
        
        # Method 2: Education section parsing
        education_patterns = [
            r'education[:\-\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*[A-Z][^:\n]*:|$)',
            r'academic.*?(?:background|qualification|record)[:\-\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*[A-Z][^:\n]*:|$)',
            r'qualification[s]?[:\-\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*[A-Z][^:\n]*:|$)'
        ]
        
        for pattern in education_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                edu_section = match.group(1)
                
                # Extract degree information from the section
                for degree_pattern in self.degree_patterns:
                    degree_matches = re.findall(degree_pattern, edu_section, re.IGNORECASE)
                    for degree_match in degree_matches:
                        if degree_match.strip():
                            education_info.add(degree_match.strip().title())
        
        # Method 3: Look for university/institution names
        university_patterns = [
            r'university\s+of\s+([^,\n]+)',
            r'([^,\n]*)\s+university',
            r'([^,\n]*)\s+institute\s+of\s+technology',
            r'([^,\n]*)\s+college\s+of\s+([^,\n]+)'
        ]
        
        for pattern in university_patterns:
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    for m in match:
                        if m.strip() and len(m.strip()) > 3:
                            education_info.add(m.strip().title())
                else:
                    if match.strip() and len(match.strip()) > 3:
                        education_info.add(match.strip().title())
        
        # Clean and filter education info
        cleaned_education = []
        for edu in education_info:
            if len(edu) > 3 and len(edu) < 100:  # Reasonable length
                cleaned_education.append(edu)
        
        return cleaned_education if cleaned_education else None

    def extract_location(self, text: str) -> Optional[str]:
        """Extract location from resume text with enhanced patterns"""
        # Enhanced location patterns
        location_patterns = [
            r'(?:address|location|city|residence|based\s+in)[:\-\s]*([^,\n]+(?:,\s*[^,\n]+)*)',
            r'\b([A-Za-z\s]+,\s*[A-Za-z\s]+,\s*\d{5,6})\b',  # City, State, PIN
            r'\b([A-Za-z\s]+,\s*[A-Za-z\s]+)\b(?=\s*\d{5,6})',  # City, State before PIN
            r'(?:from|in)\s+([A-Za-z\s]+,\s*[A-Za-z\s]+)',  # "from City, State"
            r'\b(\d{5,6})\s*,?\s*([A-Za-z\s]+(?:,\s*[A-Za-z\s]+)?)\b'  # PIN, City, State
        ]
        
        for pattern in location_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                location = matches[0]
                if isinstance(location, tuple):
                    location = ', '.join(filter(None, location))
                
                location = location.strip()
                if 3 < len(location) < 100:
                    return location
        
        # Use spaCy for location extraction if available
        if self.nlp:
            doc = self.nlp(text)
            locations = []
            for ent in doc.ents:
                if ent.label_ in ["GPE", "LOC"]:  # Geopolitical entity or location
                    locations.append(ent.text)
            
            if locations:
                return ", ".join(locations[:2])  # Return first two locations
        
        return None

    def extract_experience(self, text: str) -> Optional[float]:
        """Extract total years of experience with improved accuracy"""
        text_lower = text.lower()
        
        # Enhanced experience patterns
        experience_patterns = [
            r'(\d+(?:\.\d+)?)\s*(?:\+)?\s*years?\s+(?:of\s+)?(?:work\s+)?experience',
            r'experience.*?(\d+(?:\.\d+)?)\s*(?:\+)?\s*years?',
            r'(\d+(?:\.\d+)?)\s*(?:\+)?\s*yrs?\s+(?:of\s+)?(?:work\s+)?experience',
            r'experience.*?(\d+(?:\.\d+)?)\s*(?:\+)?\s*yrs?',
            r'(\d+(?:\.\d+)?)\s*(?:\+)?\s*years?\s+(?:in|of|with)',
            r'over\s+(\d+(?:\.\d+)?)\s*years?',
            r'more\s+than\s+(\d+(?:\.\d+)?)\s*years?',
            r'(\d+(?:\.\d+)?)\+\s*years?'
        ]
        
        years = []
        
        for pattern in experience_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                try:
                    year_value = float(match)
                    if 0 <= year_value <= 50:  # Reasonable range
                        years.append(year_value)
                except ValueError:
                    continue
        
        if years:
            return max(years)  # Return the highest experience mentioned
        
        # Try to calculate from employment dates
        date_patterns = [
            r'(\d{4})\s*[-–—]\s*(\d{4}|present|current)',
            r'(\w+\s+\d{4})\s*[-–—]\s*(\w+\s+\d{4}|present|current)',
            r'(\d{1,2}/\d{4})\s*[-–—]\s*(\d{1,2}/\d{4}|present|current)'
        ]
        
        employment_years = []
        current_year = datetime.now().year
        
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    start_str, end_str = match
                    
                    # Extract start year
                    start_year_match = re.search(r'\d{4}', start_str)
                    if not start_year_match:
                        continue
                    start_year = int(start_year_match.group())
                    
                    # Extract end year
                    if any(word in end_str.lower() for word in ['present', 'current', 'now']):
                        end_year = current_year
                    else:
                        end_year_match = re.search(r'\d{4}', end_str)
                        if not end_year_match:
                            continue
                        end_year = int(end_year_match.group())
                    
                    # Validate years
                    if 1990 <= start_year <= current_year and start_year <= end_year:
                        employment_years.append(end_year - start_year)
                
                except (ValueError, AttributeError):
                    continue
        
        return sum(employment_years) if employment_years else None

    def parse_resume(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Parse a single resume file and extract information"""
        start_time = datetime.now()
        
        try:
            # Determine file type and extract text
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                text = self.extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                text = self.extract_text_from_txt(file_path)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
            
            if not text.strip():
                logger.warning(f"No text extracted from {file_path}")
                return None
            
            # Extract information
            parsed_data = {
                'file_name': os.path.basename(file_path),
                'file_path': file_path,
                'name': self.extract_name(text),
                'email': self.extract_email(text),
                'phone_number': self.extract_phone(text),
                'skills': self.extract_skills(text),
                'education': self.extract_education(text),
                'location': self.extract_location(text),
                'total_experience': self.extract_experience(text),
                'processed_at': datetime.now().isoformat(),
                'processing_time': (datetime.now() - start_time).total_seconds()
            }
            
            # Update statistics
            self.processing_stats['total_processed'] += 1
            for field in ['name', 'email', 'phone_number', 'skills', 'education', 'location', 'total_experience']:
                if parsed_data.get(field):
                    self.processing_stats['successful_extractions'][field] += 1
            
            logger.info(f"✓ Successfully parsed {os.path.basename(file_path)}")
            return parsed_data
            
        except Exception as e:
            logger.error(f"✗ Error parsing {file_path}: {str(e)}")
            self.processing_stats['failed_files'].append({
                'file': os.path.basename(file_path),
                'error': str(e)
            })
            return None

    def parse_multiple_resumes(self, folder_path: str) -> List[Dict[str, Any]]:
        """Parse multiple resume files from a folder"""
        supported_extensions = ['.pdf', '.docx', '.txt']
        parsed_resumes = []
        
        if not os.path.exists(folder_path):
            logger.error(f"Folder {folder_path} does not exist!")
            return []
        
        # Get all supported files
        files = []
        for filename in os.listdir(folder_path):
            if os.path.splitext(filename)[1].lower() in supported_extensions:
                files.append(os.path.join(folder_path, filename))
        
        if not files:
            logger.warning(f"No supported resume files found in {folder_path}")
            return []
        
        logger.info(f"Found {len(files)} resume files to process...")
        
        # Process each file
        for i, file_path in enumerate(files, 1):
            filename = os.path.basename(file_path)
            logger.info(f"Processing {i}/{len(files)}: {filename}")
            
            try:
                parsed_data = self.parse_resume(file_path)
                if parsed_data:
                    parsed_resumes.append(parsed_data)
            except Exception as e:
                logger.error(f"Error processing {filename}: {str(e)}")
                continue
        
        # Log final statistics
        total_time = sum(resume.get('processing_time', 0) for resume in parsed_resumes)
        self.processing_stats['processing_time'] = total_time
        
        logger.info(f"Processing complete: {len(parsed_resumes)}/{len(files)} files successfully parsed")
        return parsed_resumes

    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get detailed processing statistics"""
        return {
            'total_processed': self.processing_stats['total_processed'],
            'successful_extractions': dict(self.processing_stats['successful_extractions']),
            'failed_files': self.processing_stats['failed_files'],
            'processing_time': round(self.processing_stats['processing_time'], 2),
            'success_rates': {
                field: round((count / max(self.processing_stats['total_processed'], 1)) * 100, 1)
                for field, count in self.processing_stats['successful_extractions'].items()
            } if self.processing_stats['total_processed'] > 0 else {}
        }

    def reset_statistics(self):
        """Reset processing statistics"""
        self.processing_stats = {
            'total_processed': 0,
            'successful_extractions': defaultdict(int),
            'failed_files': [],
            'processing_time': 0
        }

    def display_results(self, parsed_resumes: List[Dict[str, Any]]):
        """Display parsed results in formatted output"""
        if not parsed_resumes:
            print("No resumes were successfully parsed.")
            return
        
        print(f"\n{'='*70}")
        print(f"PARSED RESUME DATA ({len(parsed_resumes)} resumes)")
        print(f"{'='*70}\n")
        
        for i, resume_data in enumerate(parsed_resumes, 1):
            print(f"Resume {i}: {resume_data['file_name']}")
            print("-" * 50)
            
            # Display in organized format
            fields_display = [
                ('Name', resume_data.get('name', 'Not found')),
                ('Email', resume_data.get('email', 'Not found')),
                ('Phone', resume_data.get('phone_number', 'Not found')),
                ('Location', resume_data.get('location', 'Not found')),
                ('Experience', f"{resume_data.get('total_experience', 'Not found')} years" 
                             if resume_data.get('total_experience') else 'Not found'),
            ]
            
            for field_name, field_value in fields_display:
                print(f"{field_name:12}: {field_value}")
            
            # Skills
            skills = resume_data.get('skills', [])
            if skills:
                print(f"{'Skills':12}: {', '.join(skills[:10])}")  # Show first 10 skills
                if len(skills) > 10:
                    print(f"{'':12}  ... and {len(skills) - 10} more")
            else:
                print(f"{'Skills':12}: Not found")
            
            # Education
            education = resume_data.get('education', [])
            if education:
                print(f"{'Education':12}: {education[0]}")  # Show first education entry
                if len(education) > 1:
                    for edu in education[1:3]:  # Show up to 2 more
                        print(f"{'':12}  {edu}")
                    if len(education) > 3:
                        print(f"{'':12}  ... and {len(education) - 3} more")
            else:
                print(f"{'Education':12}: Not found")
            
            print(f"{'Processed':12}: {resume_data.get('processed_at', 'Unknown')}")
            print()
        
        # Display statistics
        stats = self.get_processing_statistics()
        if stats['total_processed'] > 0:
            print(f"{'='*70}")
            print("EXTRACTION STATISTICS")
            print(f"{'='*70}")
            print(f"Total files processed: {stats['total_processed']}")
            print(f"Total processing time: {stats['processing_time']} seconds")
            print(f"Average time per file: {stats['processing_time'] / stats['total_processed']:.2f} seconds")
            print()
            print("Field extraction success rates:")
            for field, rate in stats['success_rates'].items():
                count = stats['successful_extractions'][field]
                print(f"  {field.replace('_', ' ').title():15}: {count:3d}/{stats['total_processed']} ({rate:5.1f}%)")
            
            if stats['failed_files']:
                print(f"\nFailed files ({len(stats['failed_files'])}):")
                for failed in stats['failed_files']:
                    print(f"  - {failed['file']}: {failed['error']}")

    def save_to_json(self, parsed_resumes: List[Dict[str, Any]], output_file: str = 'parsed_resumes.json'):
        """Save parsed results to JSON file with metadata"""
        try:
            # Prepare data with metadata
            export_data = {
                'metadata': {
                    'exported_at': datetime.now().isoformat(),
                    'total_resumes': len(parsed_resumes),
                    'parser_version': '2.0.0',
                    'statistics': self.get_processing_statistics()
                },
                'resumes': parsed_resumes
            }
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"✓ Results saved to {output_file}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error saving to JSON: {str(e)}")
            return False

    def save_to_excel(self, parsed_resumes: List[Dict[str, Any]], output_file: str = 'parsed_resumes.xlsx'):
        """Save parsed results to Excel file"""
        try:
            if not parsed_resumes:
                logger.warning("No resumes to export to Excel.")
                return False
            df = pd.DataFrame(parsed_resumes)
            df.to_excel(output_file, index=False)
            logger.info(f"✓ Results exported to {output_file}")
            return True
        except Exception as e:
            logger.error(f"❌ Error exporting to Excel: {str(e)}")
            return False

    def validate_extracted_data(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate and clean extracted data"""
        validated_data = parsed_data.copy()
        
        # Email validation
        if validated_data.get('email'):
            email = validated_data['email']
            email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
            if not re.match(email_pattern, email):
                logger.warning(f"Invalid email format detected: {email}")
                validated_data['email'] = None
        
        # Phone validation
        if validated_data.get('phone_number'):
            phone = validated_data['phone_number']
            # Remove non-digit characters except +
            clean_phone = re.sub(r'[^\d+]', '', phone)
            if len(clean_phone.replace('+', '')) < 10:
                logger.warning(f"Invalid phone number detected: {phone}")
                validated_data['phone_number'] = None
            else:
                validated_data['phone_number'] = clean_phone
        
        # Experience validation
        if validated_data.get('total_experience'):
            exp = validated_data['total_experience']
            if not isinstance(exp, (int, float)) or exp < 0 or exp > 50:
                logger.warning(f"Invalid experience value detected: {exp}")
                validated_data['total_experience'] = None
        
        # Skills deduplication and cleaning
        if validated_data.get('skills'):
            skills = validated_data['skills']
            # Remove duplicates (case-insensitive) and empty strings
            cleaned_skills = []
            seen_skills = set()
            for skill in skills:
                skill_clean = skill.strip()
                if skill_clean and skill_clean.lower() not in seen_skills:
                    cleaned_skills.append(skill_clean)
                    seen_skills.add(skill_clean.lower())
            
            validated_data['skills'] = cleaned_skills if cleaned_skills else None
        
        # Education cleaning
        if validated_data.get('education'):
            education = validated_data['education']
            # Remove duplicates and very short entries
            cleaned_education = []
            seen_education = set()
            for edu in education:
                edu_clean = edu.strip()
                if edu_clean and len(edu_clean) > 3 and edu_clean.lower() not in seen_education:
                    cleaned_education.append(edu_clean)
                    seen_education.add(edu_clean.lower())
            
            validated_data['education'] = cleaned_education if cleaned_education else None
        
        return validated_data

    def extract_additional_info(self, text: str) -> Dict[str, Any]:
        """Extract additional information like certifications, languages, etc."""
        additional_info = {}
        
        # Extract certifications
        cert_patterns = [
            r'certifications?[:\-\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*[A-Z][^:\n]*:|$)',
            r'certified?\s+in\s+([^\n,]+)',
            r'certificate\s+in\s+([^\n,]+)'
        ]
        
        certifications = set()
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                if match.strip() and len(match.strip()) > 3:
                    certifications.add(match.strip().title())
        
        if certifications:
            additional_info['certifications'] = list(certifications)
        
        # Extract languages
        lang_patterns = [
            r'languages?[:\-\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*[A-Z][^:\n]*:|$)',
            r'fluent\s+in\s+([^\n,]+)',
            r'native\s+([^\n,]+)\s+speaker'
        ]
        
        languages = set()
        common_languages = [
            'english', 'hindi', 'spanish', 'french', 'german', 'chinese', 'japanese',
            'korean', 'arabic', 'russian', 'portuguese', 'italian', 'dutch',
            'tamil', 'telugu', 'bengali', 'marathi', 'gujarati', 'punjabi'
        ]
        
        text_lower = text.lower()
        for lang in common_languages:
            if lang in text_lower:
                languages.add(lang.title())
        
        for pattern in lang_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                if match.strip():
                    # Extract individual languages from the match
                    lang_words = re.findall(r'\b[A-Za-z]+\b', match)
                    for word in lang_words:
                        if word.lower() in common_languages:
                            languages.add(word.title())
        
        if languages:
            additional_info['languages'] = list(languages)
        
        # Extract social media/portfolio links
        url_patterns = [
            r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+',
            r'(?:https?://)?(?:www\.)?github\.com/[\w\-]+',
            r'(?:https?://)?(?:www\.)?[\w\-]+\.(?:com|org|net|io)/[\w\-/]*',
        ]
        
        urls = set()
        for pattern in url_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if match.strip():
                    urls.add(match.strip())
        
        if urls:
            additional_info['urls'] = list(urls)
        
        return additional_info

    def parse_resume_enhanced(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Enhanced resume parsing with additional information extraction"""
        # Get basic parsed data
        parsed_data = self.parse_resume(file_path)
        
        if not parsed_data:
            return None
        
        # Extract text again for additional processing
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            text = self.extract_text_from_txt(file_path)
        else:
            return parsed_data
        
        # Get additional information
        additional_info = self.extract_additional_info(text)
        parsed_data.update(additional_info)
        
        # Validate and clean data
        validated_data = self.validate_extracted_data(parsed_data)
        
        return validated_data

    def batch_process_with_progress(self, file_paths: List[str], 
                                  progress_callback: Optional[callable] = None) -> List[Dict[str, Any]]:
        """Process multiple files with progress tracking"""
        parsed_resumes = []
        total_files = len(file_paths)
        
        for i, file_path in enumerate(file_paths):
            try:
                parsed_data = self.parse_resume_enhanced(file_path)
                if parsed_data:
                    parsed_resumes.append(parsed_data)
                
                # Call progress callback if provided
                if progress_callback:
                    progress = ((i + 1) / total_files) * 100
                    progress_callback(progress, f"Processed {os.path.basename(file_path)}")
                    
            except Exception as e:
                logger.error(f"Error processing {file_path}: {str(e)}")
                if progress_callback:
                    progress = ((i + 1) / total_files) * 100
                    progress_callback(progress, f"Failed to process {os.path.basename(file_path)}")
        
        return parsed_resumes

# Utility functions for web application integration

def create_resume_parser() -> ResumeParser:
    """Factory function to create a resume parser instance"""
    return ResumeParser()

def parse_single_resume_file(file_path: str) -> Optional[Dict[str, Any]]:
    """Utility function to parse a single resume file"""
    parser = ResumeParser()
    return parser.parse_resume_enhanced(file_path)

def parse_multiple_resume_files(file_paths: List[str]) -> List[Dict[str, Any]]:
    """Utility function to parse multiple resume files"""
    parser = ResumeParser()
    return parser.batch_process_with_progress(file_paths)

def validate_file_type(filename: str) -> bool:
    """Validate if file type is supported"""
    supported_extensions = ['.pdf', '.docx', '.txt']
    return os.path.splitext(filename)[1].lower() in supported_extensions

def get_file_info(file_path: str) -> Dict[str, Any]:
    """Get basic information about a file"""
    try:
        stat = os.stat(file_path)
        return {
            'filename': os.path.basename(file_path),
            'size': stat.st_size,
            'size_mb': round(stat.st_size / (1024 * 1024), 2),
            'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'extension': os.path.splitext(file_path)[1].lower(),
            'is_supported': validate_file_type(file_path)
        }
    except Exception as e:
        return {
            'filename': os.path.basename(file_path),
            'error': str(e),
            'is_supported': False
        }

# Main execution for testing
if __name__ == "__main__":
    import argparse
    
    # Command line interface for testing
    parser_cli = argparse.ArgumentParser(description='Resume Parser - Enhanced Version')
    parser_cli.add_argument('--file', '-f', help='Single file to parse')
    parser_cli.add_argument('--folder', '-d', help='Folder containing resume files')
    parser_cli.add_argument('--output', '-o', help='Output JSON file', default='parsed_resumes.json')
    parser_cli.add_argument('--excel', '-x', help='Output Excel file', default=None)
    parser_cli.add_argument('--stats', '-s', action='store_true', help='Show detailed statistics')
    parser_cli.add_argument('--validate', '-v', action='store_true', help='Validate extracted data')
    
    args = parser_cli.parse_args()
    
    # Create parser instance
    resume_parser = ResumeParser()
    
    if args.file:
        # Parse single file
        print(f"Parsing single file: {args.file}")
        result = resume_parser.parse_resume_enhanced(args.file)
        if result:
            print(json.dumps(result, indent=2, ensure_ascii=False))
            if args.output:
                resume_parser.save_to_json([result], args.output)
            if args.excel:
                resume_parser.save_to_excel([result], args.excel)
        else:
            print("Failed to parse the file.")

    elif args.folder:
        # Parse folder
        print(f"Parsing folder: {args.folder}")
        results = resume_parser.parse_multiple_resumes(args.folder)
        if results:
            resume_parser.display_results(results)
            if args.output:
                resume_parser.save_to_json(results, args.output)
            if args.excel:
                resume_parser.save_to_excel(results, args.excel)
            if args.stats:
                stats = resume_parser.get_processing_statistics()
                print("\nDETAILED STATISTICS:")
                print(json.dumps(stats, indent=2))
        else:
            print("No resumes were successfully parsed.")
    
    else:
        print("Please provide either --file or --folder argument")
        print("Use --help for more information")